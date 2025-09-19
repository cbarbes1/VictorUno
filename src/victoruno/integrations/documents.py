"""Document processing utilities for VictorUno."""

import logging
from pathlib import Path
from typing import Optional, Dict, Any, List
import mimetypes

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False

from ..core.config import Config

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Process various document formats for the agent."""
    
    def __init__(self, config: Optional[Config] = None):
        """Initialize document processor."""
        self.config = config or Config.from_env()
        
        # Log available processors
        available_processors = []
        if PDF_AVAILABLE:
            available_processors.append("PDF")
        if DOCX_AVAILABLE:
            available_processors.append("DOCX")
        if BS4_AVAILABLE:
            available_processors.append("HTML")
        
        logger.info(f"Document processor initialized with support for: {', '.join(available_processors) or 'text files only'}")
    
    async def process_document(self, file_path: Path) -> Dict[str, Any]:
        """Process a document and extract its content."""
        try:
            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            
            # Check file size
            file_size = file_path.stat().st_size
            if file_size > self.config.max_file_size:
                raise ValueError(f"File too large: {file_size} bytes (max: {self.config.max_file_size})")
            
            # Determine file type
            file_extension = file_path.suffix.lower()
            mime_type, _ = mimetypes.guess_type(str(file_path))
            
            # Process based on file type
            if file_extension == ".pdf":
                content = await self._process_pdf(file_path)
            elif file_extension == ".docx":
                content = await self._process_docx(file_path)
            elif file_extension in [".txt", ".md"]:
                content = await self._process_text(file_path)
            elif file_extension in [".html", ".htm"]:
                content = await self._process_html(file_path)
            else:
                # Try to process as text
                try:
                    content = await self._process_text(file_path)
                except UnicodeDecodeError:
                    raise ValueError(f"Unsupported file format: {file_extension}")
            
            return {
                "filename": file_path.name,
                "file_path": str(file_path),
                "file_size": file_size,
                "file_type": file_extension,
                "mime_type": mime_type,
                "content": content,
                "word_count": len(content.split()) if content else 0,
                "char_count": len(content) if content else 0,
                "success": True,
                "message": f"Successfully processed {file_path.name}"
            }
            
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {e}")
            return {
                "filename": file_path.name if file_path else "unknown",
                "file_path": str(file_path) if file_path else "unknown",
                "file_size": 0,
                "file_type": "unknown",
                "mime_type": None,
                "content": "",
                "word_count": 0,
                "char_count": 0,
                "success": False,
                "message": f"Error processing document: {str(e)}"
            }
    
    async def _process_pdf(self, file_path: Path) -> str:
        """Process PDF file."""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 is not installed. Cannot process PDF files.")
        
        try:
            content = []
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            content.append(f"--- Page {page_num + 1} ---\n{page_text}")
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
                        content.append(f"--- Page {page_num + 1} ---\n[Error extracting text]")
            
            return "\n\n".join(content)
            
        except Exception as e:
            raise Exception(f"Error reading PDF file: {str(e)}")
    
    async def _process_docx(self, file_path: Path) -> str:
        """Process DOCX file."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is not installed. Cannot process DOCX files.")
        
        try:
            doc = Document(file_path)
            content = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                table_content = []
                for row in table.rows:
                    row_content = []
                    for cell in row.cells:
                        row_content.append(cell.text.strip())
                    table_content.append(" | ".join(row_content))
                
                if table_content:
                    content.append("\n--- Table ---")
                    content.extend(table_content)
                    content.append("--- End Table ---\n")
            
            return "\n".join(content)
            
        except Exception as e:
            raise Exception(f"Error reading DOCX file: {str(e)}")
    
    async def _process_text(self, file_path: Path) -> str:
        """Process text file."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            
            raise UnicodeDecodeError("Could not decode file with any supported encoding")
            
        except Exception as e:
            raise Exception(f"Error reading text file: {str(e)}")
    
    async def _process_html(self, file_path: Path) -> str:
        """Process HTML file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                html_content = file.read()
            
            if BS4_AVAILABLE:
                # Parse with BeautifulSoup to extract clean text
                soup = BeautifulSoup(html_content, 'html.parser')
                
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()
                
                # Get text content
                text = soup.get_text()
                
                # Clean up whitespace
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text = '\n'.join(chunk for chunk in chunks if chunk)
                
                return text
            else:
                # Fallback: return raw HTML
                logger.warning("BeautifulSoup not available. Returning raw HTML content.")
                return html_content
                
        except Exception as e:
            raise Exception(f"Error reading HTML file: {str(e)}")
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats."""
        formats = ["txt", "md"]
        
        if PDF_AVAILABLE:
            formats.append("pdf")
        if DOCX_AVAILABLE:
            formats.append("docx")
        if BS4_AVAILABLE:
            formats.extend(["html", "htm"])
        
        return formats
    
    def is_supported_format(self, file_path: Path) -> bool:
        """Check if a file format is supported."""
        extension = file_path.suffix.lower().lstrip('.')
        return extension in self.get_supported_formats()
    
    async def batch_process(self, file_paths: List[Path]) -> List[Dict[str, Any]]:
        """Process multiple documents."""
        results = []
        
        for file_path in file_paths:
            result = await self.process_document(file_path)
            results.append(result)
        
        return results
    
    async def summarize_document(self, content: str, max_length: int = 500) -> str:
        """Create a summary of document content."""
        if len(content) <= max_length:
            return content
        
        # Simple summarization: take first part, middle part, and end part
        chunk_size = max_length // 3
        
        start = content[:chunk_size]
        middle_start = len(content) // 2 - chunk_size // 2
        middle = content[middle_start:middle_start + chunk_size]
        end = content[-chunk_size:]
        
        summary = f"{start}\n\n... [content truncated] ...\n\n{middle}\n\n... [content truncated] ...\n\n{end}"
        
        return summary
    
    async def extract_keywords(self, content: str, max_keywords: int = 10) -> List[str]:
        """Extract keywords from document content."""
        import re
        from collections import Counter
        
        # Simple keyword extraction
        # Remove common stop words
        stop_words = {
            'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with',
            'by', 'is', 'are', 'was', 'were', 'be', 'been', 'have', 'has', 'had', 'do', 'does',
            'did', 'will', 'would', 'could', 'should', 'may', 'might', 'must', 'can', 'this',
            'that', 'these', 'those', 'i', 'you', 'he', 'she', 'it', 'we', 'they', 'me', 'him',
            'her', 'us', 'them', 'my', 'your', 'his', 'her', 'its', 'our', 'their'
        }
        
        # Extract words (letters only, length > 2)
        words = re.findall(r'\b[a-zA-Z]{3,}\b', content.lower())
        
        # Filter out stop words
        filtered_words = [word for word in words if word not in stop_words]
        
        # Count frequencies
        word_counts = Counter(filtered_words)
        
        # Return top keywords
        keywords = [word for word, count in word_counts.most_common(max_keywords)]
        
        return keywords